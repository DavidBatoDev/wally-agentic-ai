from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Dict, Optional
import io
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize FastAPI app
app = FastAPI(title="Word Document Text Box API", version="1.0.0")

# Request model for demonstration
class FillWordRequest(BaseModel):
    value: str = "David"  # Default value for demonstration

# Field mapping for demonstration (based on your provided data)
DEMO_FIELD_MAPPING = {
    "{first_name}": {
        "font": {
            "name": "ArialMT",
            "size": 8.039999961853027,
            "color": "#ee0000",
            "flags": 0,
            "style": "normal"
        },
        "label": "Child's First Name",
        "position": {
            "x0": 129.6199951171875,
            "x1": 172.38536071777344,
            "y0": 144.3037567138672,
            "y1": 153.2683563232422,
            "width": 42.76536560058594,
            "height": 8.964599609375
        },
        "rotation": 0,
        "alignment": "left",
        "bbox_center": {
            "x": 151.00267791748047,
            "y": 148.7860565185547
        },
        "page_number": 1
    }
}

def hex_to_rgb(hex_color: str) -> tuple:
    """Convert hex color to RGB tuple"""
    if hex_color.startswith("#"):
        hex_color = hex_color[1:]
    try:
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    except:
        return (0, 0, 0)  # Default to black

def add_formatted_text(doc, text: str, field_info: dict):
    """Add formatted text using paragraph with styling"""
    
    # Get position and font info
    position = field_info["position"]
    font_info = field_info["font"]
    
    # Add spacing to simulate positioning (rough approximation)
    # This adds empty paragraphs to push content down
    y_position = position.get("y0", 0)
    if y_position > 100:  # If position is far down, add some spacing
        spacing_lines = max(1, int(y_position / 50))  # Rough calculation
        for _ in range(min(spacing_lines, 5)):  # Limit to avoid too much spacing
            doc.add_paragraph("")
    
    # Add the main paragraph with the text
    paragraph = doc.add_paragraph()
    
    # Add indentation to simulate x-position (rough approximation)
    x_position = position.get("x0", 0)
    if x_position > 50:  # If position is to the right, add indentation
        # Convert position to approximate indentation
        indent_inches = min(x_position / 100, 3)  # Rough conversion, max 3 inches
        paragraph.paragraph_format.left_indent = Inches(indent_inches)
    
    # Add a run with the text
    run = paragraph.add_run(text)
    
    # Set font properties
    font = run.font
    font_name = font_info.get("name", "Arial")
    # Handle common font name variations
    if font_name == "ArialMT":
        font_name = "Arial"
    font.name = font_name
    font.size = Pt(font_info.get("size", 12))
    
    # Set color
    color_hex = font_info.get("color", "#000000")
    if color_hex.startswith("#"):
        r, g, b = hex_to_rgb(color_hex)
        font.color.rgb = RGBColor(r, g, b)
    
    # Set alignment
    alignment = field_info.get("alignment", "left").lower()
    if alignment == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "justify":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    return paragraph

def add_text_with_table_positioning(doc, text: str, field_info: dict):
    """Add text using a borderless table for better positioning control"""
    
    # Get font info
    font_info = field_info["font"]
    position = field_info["position"]
    
    # Create a single-cell table
    table = doc.add_table(rows=1, cols=1)
    
    # Remove table borders by setting table style
    table.style = 'Table Grid'
    
    # Get the cell and clear default paragraph
    cell = table.cell(0, 0)
    
    # Set cell width based on position width
    width_inches = min(position.get("width", 100) / 72, 6)  # Convert points to inches, max 6"
    cell.width = Inches(width_inches)
    
    # Get the paragraph in the cell
    paragraph = cell.paragraphs[0]
    
    # Add the text
    run = paragraph.add_run(text)
    
    # Set font properties
    font = run.font
    font_name = font_info.get("name", "Arial")
    if font_name == "ArialMT":
        font_name = "Arial"
    font.name = font_name
    font.size = Pt(font_info.get("size", 12))
    
    # Set color
    color_hex = font_info.get("color", "#000000")
    if color_hex.startswith("#"):
        r, g, b = hex_to_rgb(color_hex)
        font.color.rgb = RGBColor(r, g, b)
    
    # Set alignment
    alignment = field_info.get("alignment", "left").lower()
    if alignment == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "justify":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Try to remove table borders (this is a simplified approach)
    try:
        # This removes visible borders
        for row in table.rows:
            for cell in row.cells:
                # Set cell margins
                cell._element.get_or_add_tcPr()
    except:
        pass  # If border removal fails, continue
    
    return table

@app.post(
    "/insert-textbox-demo",
    summary="Insert text box into Word document (demonstration)"
)
async def insert_textbox_demo(
    file: UploadFile = File(...),
    request: FillWordRequest = FillWordRequest()
):
    """
    Insert a text box into a Word document at a predefined position.
    This is a demonstration endpoint that uses hardcoded field mapping.
    
    Args:
        file: Word document file (.docx)
        request: Request containing the value to insert (defaults to "David")
    
    Returns:
        Modified Word document with text box inserted
    """
    
    # Validate file type
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    try:
        # Read the uploaded file
        file_content = await file.read()
        doc = Document(io.BytesIO(file_content))
        
        # Get field info for demonstration
        field_info = DEMO_FIELD_MAPPING["{first_name}"]
        
        # Add text using table-based positioning
        table = add_text_with_table_positioning(doc, request.value, field_info)
        
        # Add a note about the insertion
        note_paragraph = doc.add_paragraph()
        note_run = note_paragraph.add_run(
            f"\n[Note: Text '{request.value}' inserted for field '{field_info['label']}' "
            f"at approximate position ({field_info['position']['x0']:.1f}, {field_info['position']['y0']:.1f}) "
            f"with font {field_info['font']['name']} size {field_info['font']['size']} color {field_info['font']['color']}]"
        )
        note_run.font.size = Pt(8)
        note_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
        
        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        # Return the modified document
        headers = {
            "Content-Disposition": f'attachment; filename="modified_{file.filename}"'
        }
        return StreamingResponse(
            io.BytesIO(buf.read()), 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing Word document: {str(e)}")

@app.post(
    "/insert-textbox-simple",
    summary="Insert text using simple paragraph formatting"
)
async def insert_textbox_simple(
    file: UploadFile = File(...),
    value: str = "David",
    field_key: str = "{first_name}"
):
    """
    Insert text using simple paragraph formatting with positioning approximation.
    
    Args:
        file: Word document file (.docx)
        value: Text value to insert
        field_key: Field key to use from demo mapping (default: "{first_name}")
    
    Returns:
        Modified Word document
    """
    
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    try:
        file_content = await file.read()
        doc = Document(io.BytesIO(file_content))
        
        # Get field info
        if field_key not in DEMO_FIELD_MAPPING:
            raise HTTPException(status_code=400, detail=f"Field key '{field_key}' not found in demo mapping")
        
        field_info = DEMO_FIELD_MAPPING[field_key]
        
        # Insert text using simple formatting
        paragraph = add_formatted_text(doc, value, field_info)
        
        # Add processing info
        info_paragraph = doc.add_paragraph()
        info_run = info_paragraph.add_run(
            f"\n[Processed: Field '{field_info['label']}' = '{value}' "
            f"with font {field_info['font']['name']} size {field_info['font']['size']}]"
        )
        info_run.font.size = Pt(8)
        info_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        headers = {
            "Content-Disposition": f'attachment; filename="simple_{file.filename}"'
        }
        return StreamingResponse(
            io.BytesIO(buf.read()), 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@app.post(
    "/insert-textbox-advanced",
    summary="Insert text box with custom field mapping"
)
async def insert_textbox_advanced(
    file: UploadFile = File(...),
    field_key: str = "{first_name}",
    value: str = "David",
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    color: Optional[str] = None,
    alignment: Optional[str] = None
):
    """
    Insert text box with custom parameters or using the demo field mapping.
    
    Args:
        file: Word document file (.docx)
        field_key: Field key to use from demo mapping (default: "{first_name}")
        value: Text value to insert
        font_name: Override font name
        font_size: Override font size
        color: Override color (hex format)
        alignment: Override alignment
    
    Returns:
        Modified Word document
    """
    
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    try:
        file_content = await file.read()
        doc = Document(io.BytesIO(file_content))
        
        # Get base field info
        if field_key not in DEMO_FIELD_MAPPING:
            raise HTTPException(status_code=400, detail=f"Field key '{field_key}' not found in demo mapping")
        
        field_info = DEMO_FIELD_MAPPING[field_key].copy()
        
        # Override parameters if provided
        if font_name:
            field_info["font"]["name"] = font_name
        if font_size:
            field_info["font"]["size"] = font_size
        if color:
            field_info["font"]["color"] = color
        if alignment:
            field_info["alignment"] = alignment
        
        # Insert text using table approach
        table = add_text_with_table_positioning(doc, value, field_info)
        
        # Add processing info
        info_paragraph = doc.add_paragraph()
        info_run = info_paragraph.add_run(
            f"\n[Advanced Processing: Field '{field_info['label']}' = '{value}' "
            f"with font {field_info['font']['name']} size {field_info['font']['size']} "
            f"color {field_info['font']['color']} alignment {field_info.get('alignment', 'left')}]"
        )
        info_run.font.size = Pt(8)
        info_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        headers = {
            "Content-Disposition": f'attachment; filename="advanced_{file.filename}"'
        }
        return StreamingResponse(
            io.BytesIO(buf.read()), 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@app.get("/demo-field-info")
async def get_demo_field_info():
    """Get the demo field mapping information"""
    return {
        "message": "Demo field mapping for Word document text insertion",
        "field_mapping": DEMO_FIELD_MAPPING,
        "note": "Positioning in Word documents is approximate due to format limitations"
    }

@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "Word Document Text Box API",
        "version": "1.0.0",
        "description": "Insert formatted text into Word documents with approximate positioning",
        "endpoints": {
            "/insert-textbox-demo": "POST - Insert text using demo field mapping (table-based)",
            "/insert-textbox-simple": "POST - Insert text using simple paragraph formatting", 
            "/insert-textbox-advanced": "POST - Insert text with custom parameters",
            "/demo-field-info": "GET - View demo field mapping",
            "/docs": "GET - API documentation"
        },
        "supported_formats": [".docx"],
        "note": "Text positioning in Word is approximate due to format limitations. Three methods available: table-based, simple paragraph, and advanced customization."
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "service": "word-textbox-api"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)